# -*- coding: utf-8 -*-
"""
Generador Maestro Definitivo de Informes Sprint 1 (USMP / APA 7ma Edicion).
Asegura:
1. Portada oficial USMP en exactamente 1 pagina (P[00] a P[37] con line_spacing=1.0, space_before=0, space_after=0).
2. Salto de pagina limpio en P[38] (Title).
3. 100% UTF-8 limpio con todas las tildes, acentos, letras enie y guiones ortograficos en espanol correctos.
4. Redaccion tecnica formal, profesional y academica de nivel universitario.
5. Formato de figuras APA 7 (Caption 'Figura X', Titulo en cursiva, Imagen centrada 5.2 pulg, Nota 'Nota. ... Fuente: ...' en 10pt).
6. Referencias con sangria francesa (1.27 cm).
7. Encabezado con numero de pagina y Pie de pagina oficial con autor y rol.
"""

import shutil
import tempfile
import pathlib
import time
import comtypes.client as cc
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Rutas base
S0_DIR = pathlib.Path(r"C:\Users\gabri\OneDrive\Desktop\AgroSmart-Insights-Informes\Entrega Sprint 0")
TEMPLATE_PATH = S0_DIR / "Informe-03-Configuracion-GitHub-Sprint0-V.3.docx"
LOGO_PATH = S0_DIR / "assets" / "logo-usmp.png"

OUT_DIR = pathlib.Path(r"C:\Users\gabri\OneDrive\Desktop\AgroSmart-Insights-Informes\Entrega Sprint 1")
EVID_DIR = OUT_DIR / "assets" / "evidencias"
ASSETS_S1_DIR = OUT_DIR / "assets"
REPO_DOCS_DIR = pathlib.Path(r"C:\Users\gabri\OneDrive\Documentos\Agrosmart-Insights\docs\sprint-1\informes\devsecops")

# Metadatos del autor y contexto academico
INTEGRANTE = "León Cangalaya, Gabriel Emilio"
ROL = "Líder DevSecOps – Célula 2"
FECHA = "25/08/2026"
CURSO = "TALLER DE PROYECTOS"
PROFESOR = "ING. NORMA VIRGINIA LEÓN LESCANO"
PROYECTO = "AGROSMART INSIGHTS – OPTIMIZACIÓN DE MERCADOS Y PRECIOS AGRÍCOLAS"
FACULTAD = "ESCUELA PROFESIONAL DE INGENIERÍA DE COMPUTACIÓN Y SISTEMAS"
CIUDAD = "LIMA – PERÚ"


def init_document(template_path=TEMPLATE_PATH):
    """Crea un nuevo Document a partir de la plantilla Sprint 0 preservando sectPr, header y footer."""
    tmp = pathlib.Path(tempfile.mktemp(suffix='.docx'))
    shutil.copy2(template_path, tmp)
    doc = Document(str(tmp))
    
    # Limpiar todos los elementos del cuerpo excepto la definicion de seccion (sectPr)
    for el in list(doc.element.body):
        if el.tag.endswith('sectPr'):
            continue
        doc.element.body.remove(el)
    return doc


def apply_p_spacing(paragraph, line_spacing=240, before=0, after=0):
    """Aplica espaciado XML estricto a un parrafo."""
    pPr = paragraph._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'), str(line_spacing))
    sp.set(qn('w:lineRule'), 'auto')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    pPr.append(sp)


def build_portada(doc, titulo_informe, subtitulo_informe=None, fecha=FECHA):
    """
    Construye la portada formal USMP exactamente como en el Sprint 0.
    Garantiza que toda la portada quede contenida en la Pagina 1 sin desbordamiento.
    """
    def add_port_p(text="", size=None, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_p_spacing(p, line_spacing=240, before=0, after=0)
        if text:
            r = p.add_run(text)
            r.bold = bold
            if size:
                r.font.size = Pt(size)
        return p

    # P[00] Logo oficial de la USMP centrado
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_p_spacing(p0, line_spacing=240, before=0, after=0)
    r0 = p0.add_run()
    r0.add_picture(str(LOGO_PATH), width=Inches(4.33))

    # P[01] Espacio
    add_port_p()

    # P[02] Escuela Profesional
    add_port_p(FACULTAD, size=11.0, bold=True)

    # P[03]..P[07] Espacios
    for _ in range(5):
        add_port_p()

    # P[08] Titulo del Informe
    add_port_p(titulo_informe, size=14.0, bold=True)

    if subtitulo_informe:
        add_port_p(subtitulo_informe, size=12.0, bold=True)
        # Ajustar espacios
        num_espacios = 4
    else:
        num_espacios = 5

    # P[09]..P[13] Espacios
    for _ in range(num_espacios):
        add_port_p()

    # P[14] Etiqueta Proyecto
    add_port_p("PROYECTO:", bold=True)

    # P[15] Nombre del Proyecto
    add_port_p(PROYECTO, bold=True)

    # P[16]..P[17] Espacios
    for _ in range(2):
        add_port_p()

    # P[18] Etiqueta Curso
    add_port_p("CURSO:", bold=True)

    # P[19] Nombre del Curso
    add_port_p(CURSO, bold=True)

    # P[20]..P[21] Espacios
    for _ in range(2):
        add_port_p()

    # P[22] Etiqueta Profesor
    add_port_p("PROFESOR:", bold=True)

    # P[23] Nombre del Profesor
    add_port_p(PROFESOR, bold=True)

    # P[24]..P[25] Espacios
    for _ in range(2):
        add_port_p()

    # P[26] Etiqueta Integrante
    add_port_p("INTEGRANTE:", bold=True)

    # P[27] Nombre del Integrante
    add_port_p(INTEGRANTE, bold=True)

    # P[28] Rol
    add_port_p(ROL, bold=True)

    # P[29]..P[31] Espacios
    for _ in range(3):
        add_port_p()

    # P[32] Fecha
    add_port_p(f"FECHA: {fecha}", bold=True)

    # P[33]..P[36] Espacios
    for _ in range(4):
        add_port_p()

    # P[37] Ciudad
    add_port_p(CIUDAD, bold=True)


def add_doc_title(doc, text):
    """Inserta el titulo del documento en la Pagina 2 mediante pageBreakBefore."""
    p = doc.add_paragraph(style='Title')
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pPr.append(pb)
    p.add_run(text)
    return p


def add_heading_1(doc, text):
    return doc.add_heading(text, level=1)


def add_heading_2(doc, text):
    return doc.add_heading(text, level=2)


def add_paragraph_body(doc, text):
    """Inserta parrafo de cuerpo normal con sangria de primera linea APA de 1.27 cm."""
    p = doc.add_paragraph(text, style='Normal')
    p.paragraph_format.first_line_indent = Cm(1.27)
    return p


def add_bullet_item(doc, text):
    """Inserta item de lista con viñeta."""
    return doc.add_paragraph(text, style='List Bullet')


def add_figure_apa7(doc, figure_num, title, img_path, note_text, source_text="Elaboración propia a partir del repositorio Agrosmart-Insights (Equipo 2 AgroSmart Insights, 2026)."):
    """
    Inserta una figura con formato estricto APA 7ma edicion:
    1. Etiqueta Caption: Figura X
    2. Titulo descriptivo en cursiva
    3. Imagen centrada
    4. Nota al pie con 'Nota. [texto] Fuente: [texto]' en tamaño 10 pt
    """
    # 1. Caption
    p_cap = doc.add_paragraph(style='Caption')
    p_cap.add_run(f"Figura {figure_num}")

    # 2. Titulo de la figura en cursiva
    p_title = doc.add_paragraph(style='Normal')
    p_title.paragraph_format.first_line_indent = Cm(0)
    p_title.add_run(title).italic = True

    # 3. Imagen centrada
    path_obj = pathlib.Path(img_path)
    if path_obj.exists():
        p_img = doc.add_paragraph(style='Normal')
        p_img.paragraph_format.first_line_indent = Cm(0)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        r_img.add_picture(str(path_obj), width=Inches(5.2))

    # 4. Nota explicativa en 10 pt
    p_note = doc.add_paragraph(style='Normal')
    p_note.paragraph_format.first_line_indent = Cm(0)
    
    r_nota_label = p_note.add_run("Nota. ")
    r_nota_label.italic = True
    r_nota_label.font.size = Pt(10.0)

    r_nota_desc = p_note.add_run(f"{note_text} ")
    r_nota_desc.font.size = Pt(10.0)

    r_fuente_label = p_note.add_run("Fuente: ")
    r_fuente_label.italic = True
    r_fuente_label.font.size = Pt(10.0)

    r_fuente_text = p_note.add_run(source_text)
    r_fuente_text.font.size = Pt(10.0)


def add_references_section(doc, references_list):
    """Inserta seccion de referencias con sangria francesa de 1.27 cm."""
    add_heading_1(doc, "Referencias")
    for ref_text in references_list:
        p = doc.add_paragraph(ref_text, style='Normal')
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)


# ==============================================================================
# 1. GENERACIÓN: INFORME DE CONFIGURACIÓN DE VERCEL
# ==============================================================================
def generate_vercel_report():
    print("\n--- Generando Informe de Configuración de Vercel ---")
    doc = init_document()
    build_portada(doc, "INFORME DE CONFIGURACIÓN DE VERCEL")

    add_doc_title(doc, "INFORME DE CONFIGURACIÓN DE VERCEL – AGROSMART INSIGHTS")

    add_heading_1(doc, "1. Objetivo")
    add_paragraph_body(doc, "El presente informe tiene por objetivo documentar, sustentar y verificar la configuración del despliegue del frontend de AgroSmart Insights en la plataforma Vercel durante el Sprint 1. El alcance técnico abarca la importación del repositorio oficial desde GitHub, la delimitación del subdirectorio raíz (Root Directory) en una arquitectura monorepo, la validación del archivo de configuración vercel.json, la parametrización de variables de entorno para el enlace con el backend n8n en Render y la comprobación de la integración continua.")

    add_heading_2(doc, "1.1. Fundamento metodológico y enfoque DevSecOps")
    add_paragraph_body(doc, "La arquitectura de AgroSmart Insights implementa un desacoplamiento estricto entre la interfaz de usuario desarrollada en React con Vite y los servicios de orquestación analítica e inteligencia artificial alojados en Render. Este modelo arquitectónico demanda que la capa frontend sea servida a través de una red perimetral de distribución global (Edge Network), lo cual optimiza los tiempos de carga, minimiza la latencia de respuesta y reduce la superficie de ataque al no exponer servidores de aplicación directos (Vercel, 2026).")
    add_paragraph_body(doc, "Desde la perspectiva de gobernanza y seguridad DevSecOps, se aplican los principios de entrega segura del marco NIST SP 800-218. Cada modificación introducida en la interfaz gráfica debe ser evaluada automáticamente mediante el pipeline validate-frontend en GitHub Actions antes de ser admitida en la rama main, impidiendo la introducción de vulnerabilidades de código, malas prácticas de sintaxis o filtración accidental de credenciales (NIST, 2022).")

    add_heading_2(doc, "1.2. Arquitectura de despliegue frontend")
    add_paragraph_body(doc, "El módulo frontend agrosmart-frontend ofrece una interfaz de lenguaje natural (NLQ) orientada a productores agrícolas y comerciantes del Gran Mercado Mayorista de Lima (GMML). Las peticiones emitidas por los usuarios son procesadas por el cliente asíncrono n8nClient.js y transmitidas mediante peticiones HTTP POST seguras hacia el endpoint webhook del workflow WF2 desplegado en Render.")

    add_figure_apa7(
        doc, 1,
        "Arquitectura de comunicación desacoplada entre Frontend Vercel y Backend Render",
        ASSETS_S1_DIR / "fig_arquitectura-demo-sprint1.png",
        "Diagrama que ilustra la separación funcional de capas. Los clientes acceden a la aplicación React servida por la red Edge de Vercel, la cual se comunica asíncronamente con el webhook expuesto por n8n sobre PostgreSQL en Render.",
        "Elaboración propia a partir del diseño de arquitectura de AgroSmart Insights (Equipo 2 AgroSmart Insights, 2026)."
    )

    add_heading_1(doc, "2. Procedimiento de configuración en Vercel")

    add_heading_2(doc, "2.1. Conexión e importación del repositorio")
    add_paragraph_body(doc, "Se inició sesión en la consola de administración de Vercel (https://vercel.com) vinculada con la cuenta de la organización institucional en GitHub. Se seleccionó la opción Add New → Project y se localizó el repositorio equipo2-AgroSmart-Insights/Agrosmart-Insights, estableciendo main como la rama de producción predeterminada.")

    add_figure_apa7(
        doc, 2,
        "Importación del repositorio Agrosmart-Insights en la plataforma Vercel",
        EVID_DIR / "fig01-vercel-importar-proyecto.png",
        "Captura de pantalla de la consola de Vercel en la sección Add New Project. Se visualiza la selección del repositorio institucional, la rama de producción main y el nombre del proyecto agrosmart-insights.",
        "Elaboración propia a partir de la consola de Vercel (Vercel, 2026)."
    )

    add_heading_2(doc, "2.2. Configuración del Root Directory en entorno Monorepo")
    add_paragraph_body(doc, "El repositorio de AgroSmart Insights está estructurado como un monorepo que alberga simultáneamente el código del frontend, flujos de trabajo de n8n, migraciones de base de datos y scripts de automatización. Debido a que la aplicación React no se localiza en la raíz del repositorio (./), fue necesario ajustar el parámetro Root Directory.")
    add_paragraph_body(doc, "Utilizando el selector interactivo de directorios de Vercel, se navegó hacia el subdirectorio src/frontend/agrosmart-frontend. Al confirmar dicha ubicación, el motor de análisis de Vercel detectó automáticamente el archivo package.json y configuró el entorno de compilación optimizado para Vite.")

    add_figure_apa7(
        doc, 3,
        "Selección del subdirectorio src/frontend/agrosmart-frontend como Root Directory",
        EVID_DIR / "fig02-vercel-root-directory.png",
        "Captura de pantalla del selector de directorio raíz de Vercel. Se observa la navegación hacia src/frontend/agrosmart-frontend y el reconocimiento automático del framework Vite mediante su ícono distintivo.",
        "Elaboración propia a partir de la consola de Vercel (Vercel, 2026)."
    )

    add_heading_2(doc, "2.3. Parámetros de compilación y empaquetado (Build Settings)")
    add_paragraph_body(doc, "Se validó la coherencia de los comandos de construcción con el archivo vercel.json presente en el frontend. Los parámetros configurados son los siguientes:")
    add_bullet_item(doc, "Framework Preset: Vite (reconocimiento automático).")
    add_bullet_item(doc, "Build Command: npm run build (genera los artefactos estáticos minificados en el directorio dist/).")
    add_bullet_item(doc, "Output Directory: dist (directorio público distribuido por la red de Vercel).")
    add_bullet_item(doc, "Install Command: npm ci (garantiza una instalación determinista basada en package-lock.json).")

    add_heading_2(doc, "2.4. Configuración de Variables de Entorno")
    add_paragraph_body(doc, "Para habilitar la conectividad en tiempo de ejecución entre la interfaz gráfica y el motor analítico de n8n, se registró la variable de entorno en la pantalla New Project y, de forma persistente, en Project Settings → Environment Variables:")
    add_bullet_item(doc, "VITE_N8N_WEBHOOK_URL: URL pública absoluta del webhook de producción en Render (https://agrosmart-n8n.onrender.com/webhook/v1/query).")
    add_bullet_item(doc, "Ámbitos: Production y Preview, de modo que las vistas previas de rama también apunten al backend cloud.")
    add_paragraph_body(doc, "El prefijo VITE_ permite que el empaquetador inyecte el valor durante la fase de compilación sin comprometer secretos sensibles del servidor. Las llaves de Groq, Gemini, Hugging Face y MapTiler permanecen exclusivamente en Render.")

    add_figure_apa7(
        doc, 4,
        "Variable de entorno VITE_N8N_WEBHOOK_URL en el despliegue de Vercel",
        EVID_DIR / "fig15-vercel-env-webhook.png",
        "Captura de la pantalla New Project de Vercel. Se observa el framework Vite, la variable VITE_N8N_WEBHOOK_URL enmascarada y los ámbitos Production y Preview. El Root Directory debe quedar en src/frontend/agrosmart-frontend antes de pulsar Deploy.",
        "Elaboración propia a partir de la consola de Vercel (Vercel, 2026)."
    )

    add_heading_1(doc, "3. Verificación, Pipeline CI/CD y Despliegue")

    add_heading_2(doc, "3.1. Integración con el pipeline GitHub Actions")
    add_paragraph_body(doc, "En cumplimiento con las políticas de control de calidad, el flujo de trabajo frontend-ci.yml ejecuta el análisis estático de código (ESLint) y la prueba de empaquetado (npm run build) de forma obligatoria ante cada Pull Request dirigido a main. Asimismo, Vercel genera despliegues de vista previa (Preview Deployments) por cada rama secundaria para facilitar la revisión visual.")

    add_heading_2(doc, "3.2. Verificación de despliegue en producción")
    add_paragraph_body(doc, "El 25 de agosto de 2026 se ejecutó el Deploy inicial del proyecto agrosmart-insights. Vercel completó la compilación Vite y mostró la pantalla Congratulations con vista previa de la interfaz AgroSmart Insights, confirmando el estado Ready y el cifrado SSL/TLS automático.")

    add_figure_apa7(
        doc, 5,
        "Despliegue exitoso del frontend AgroSmart Insights en Vercel",
        EVID_DIR / "fig16-vercel-deploy-ok.png",
        "Captura de la pantalla Congratulations de Vercel. Se visualiza la vista previa de AgroSmart Insights ya publicada en la red Edge.",
        "Elaboración propia a partir de la consola de Vercel (Vercel, 2026)."
    )

    add_heading_2(doc, "3.3. Prueba de extremo a extremo y ajuste de timeout")
    add_paragraph_body(doc, "Con el webhook de n8n ya registrado en producción, se ejecutó una consulta de prueba en el chat: «¿Cuál es el precio actual de la papa?». El frontend abortó a los 6 segundos porque n8nClient.js tenía TIMEOUT_MS = 6000. El flujo WF2 en Render (clasificación LLM, consulta PostgreSQL y redacción) tarda habitualmente entre 8 y 20 segundos en el plan gratuito, de modo que el corte no era un fallo de n8n sino un límite del cliente.")
    add_paragraph_body(doc, "Se actualizó el cliente a un timeout de 30 segundos (configurable con VITE_API_TIMEOUT_MS) y se documentó el valor en .env.example. Este ajuste forma parte del presente Pull Request; hasta que se fusione en main, la prueba E2E completa en Vercel seguirá limitada por los 6 segundos del bundle publicado.")

    add_figure_apa7(
        doc, 6,
        "Error de timeout de 6 segundos en el chat publicado en Vercel",
        EVID_DIR / "fig19-frontend-timeout-6s.png",
        "Captura de AgroSmart Insights en producción. La consulta se procesó en n8n, pero el navegador canceló la espera al cumplir el límite de 6 segundos del cliente.",
        "Elaboración propia a partir de la aplicación desplegada en Vercel."
    )

    add_heading_1(doc, "4. Consideraciones de Seguridad DevSecOps")
    add_paragraph_body(doc, "El despliegue en Vercel incorpora las siguientes medidas de protección:")
    add_bullet_item(doc, "Aislamiento de claves de IA: Ninguna llave privada (Groq, Gemini, Hugging Face) reside en el bundle del cliente; todas se conservan aisladas en el backend de Render.")
    add_bullet_item(doc, "Inmutabilidad de artefactos: Cada despliegue se encuentra asociado al hash de commit criptográfico inalterable de GitHub.")
    add_bullet_item(doc, "Tránsito cifrado obligatorio: Se implementa HTTPS estricto con cabeceras HSTS para toda la comunicación web.")

    add_heading_1(doc, "5. Conclusión")
    add_paragraph_body(doc, "La puesta en marcha del frontend de AgroSmart Insights en Vercel quedó evidenciada el 25 de agosto de 2026: repositorio importado, Root Directory del monorepo, variable VITE_N8N_WEBHOOK_URL apuntando a Render y build Ready. El único refinamiento pendiente de fusión es el timeout de 30 segundos del cliente NLQ, necesario para que las respuestas de WF2 no se corten en el plan gratuito de Render.")

    add_references_section(doc, [
        "Equipo 2 AgroSmart Insights. (2026). Agrosmart-Insights [Repositorio de software]. GitHub. https://github.com/equipo2-AgroSmart-Insights/Agrosmart-Insights",
        "GitHub. (s. f.). Acerca de las ramas protegidas. Documentación de GitHub. https://docs.github.com/es/repositories/configuring-branches-and-merges-in-your-repository/defining-the-mergeability-of-pull-requests/about-protected-branches",
        "National Institute of Standards and Technology. (2022). Secure Software Development Framework (SSDF) Version 1.1 (NIST SP 800-218). https://csrc.nist.gov/Projects/ssdf",
        "Red Hat. (s. f.). ¿Qué es DevSecOps? https://www.redhat.com/es/topics/devops/what-is-devsecops",
        "Render. (2026). Blueprint specification. https://render.com/docs/blueprint-spec",
        "Vercel. (2026). Deploying a Git repository. https://vercel.com/docs/deployments/git",
        "Vercel. (2026). Monorepos. https://vercel.com/docs/monorepos",
    ])

    out_docx = OUT_DIR / "Informe-Configuracion-Vercel-Sprint1-V.1.docx"
    doc.save(out_docx)
    print(f"-> Guardado exitosamente: {out_docx}")
    return out_docx


# ==============================================================================
# 2. GENERACIÓN: INFORME DE CONFIGURACIÓN DE RENDER
# ==============================================================================
def generate_render_report():
    print("\n--- Generando Informe de Configuración de Render ---")
    doc = init_document()
    build_portada(doc, "INFORME DE CONFIGURACIÓN DE RENDER")

    add_doc_title(doc, "INFORME DE CONFIGURACIÓN DE RENDER – AGROSMART INSIGHTS")

    add_heading_1(doc, "1. Objetivo")
    add_paragraph_body(doc, "El propósito de este informe es documentar la arquitectura, la especificación de infraestructura como código (IaC), el procedimiento de despliegue coordinado, el diagnóstico exhaustivo de fallos y la resolución de incidencias en la plataforma Cloud Render durante el Sprint 1. Se detalla el aprovisionamiento de la base de datos PostgreSQL y del motor de automatización n8n mediante render.yaml, los ajustes de rendimiento de memoria, la configuración de reglas de red y la ejecución de migraciones con 9 445 registros de prueba.")

    add_heading_2(doc, "1.1. Fundamento metodológico")
    add_paragraph_body(doc, "Render ofrece una plataforma en la nube basada en infraestructura declarativa mediante especificaciones Blueprint (render.yaml). Este enfoque permiteversionar la infraestructura junto con el código fuente, garantizando entornos reproducibles, trazabilidad de cambios y mitigación de errores de configuración manual entre ambientes (Render, 2026).")
    add_paragraph_body(doc, "En apego a las directrices DevSecOps, la capa backend interconecta el orquestador n8n y la base de datos relacional PostgreSQL a través de una red privada virtual (Private Network), bloqueando accesos no autorizados desde internet y gestionando las credenciales y llaves de modelos de lenguaje mediante variables seguras inyectadas en tiempo de despliegue (NIST, 2022).")

    add_heading_2(doc, "1.2. Arquitectura de Backend en Render")
    add_paragraph_body(doc, "El backend de AgroSmart Insights está compuesto por dos servicios fundamentales:")
    add_bullet_item(doc, "agrosmart-postgres: Instancia gestionada de PostgreSQL 18 con la extensión pgvector habilitada, destinada al almacenamiento de catálogos maestros de productos, registros de mercados, series de precios históricos del GMML y vectores de similitud semántica para el agente RAG.")
    add_bullet_item(doc, "agrosmart-n8n: Servidor de flujos de trabajo basado en Node.js, desplegado con la versión optimizada n8n:1.83.2, responsable de la recepción de consultas en lenguaje natural, invocación de modelos LLM (Groq y Gemini) y generación de respuestas estructuradas.")

    add_figure_apa7(
        doc, 1,
        "Diagrama de arquitectura del backend DevSecOps en Render",
        ASSETS_S1_DIR / "fig_arquitectura-devsecops.png",
        "Diagrama que representa la topología interna en Render. El servicio web n8n y la base de datos PostgreSQL se comunican a través de la red privada interna con aislamiento de credenciales y exposición controlada de endpoints webhook.",
        "Elaboración propia a partir del diseño de arquitectura de AgroSmart Insights (Equipo 2 AgroSmart Insights, 2026)."
    )

    add_heading_1(doc, "2. Especificación de Infraestructura como Código (Blueprint)")

    add_heading_2(doc, "2.1. Estructura del archivo render.yaml")
    add_paragraph_body(doc, "Se diseñó el archivo render.yaml en la raíz del repositorio para definir la infraestructura de manera reproducible. La configuración incluye:")
    add_bullet_item(doc, "Servicio de Base de Datos (databases): Nombre agrosmart-postgres, base de datos agrosmart, usuario agrosmart, plan free.")
    add_bullet_item(doc, "Servicio Web de Aplicación (services): Tipo web, entorno de contenedor Docker, imagen docker.io/n8nio/n8n:1.83.2, puerto 5678 y disco de almacenamiento persistente de 1 GB montado en /home/node/.n8n.")

    add_heading_2(doc, "2.2. Parametrización y variables de entorno del Blueprint")
    add_paragraph_body(doc, "Se establecieron variables de conexión asociadas automáticamente mediante la cláusula fromDatabase (DB_POSTGRESDB_HOST, DB_POSTGRESDB_USER, DB_POSTGRESDB_PASSWORD, DB_POSTGRESDB_DATABASE), junto con variables marcadas sync: false para la captura segura de API keys en la consola de Render.")

    add_figure_apa7(
        doc, 2,
        "Conexión del repositorio en Render para inicializar el Blueprint",
        EVID_DIR / "fig03-render-conectar-repo.png",
        "Captura de pantalla de la consola de Render en el módulo Blueprints, mostrando la vinculación con el repositorio equipo2-AgroSmart-Insights/Agrosmart-Insights.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 3,
        "Formulario de configuración de variables de entorno y API Keys en Render Blueprint",
        EVID_DIR / "fig04-render-blueprint-claves.png",
        "Captura de pantalla de la interfaz de Render mostrando los campos de variables de entorno sensibles (GROQ_API_KEY, GEMINI_API_KEY, HUGGINGFACE_API_KEY, MAPTILER_API_KEY) debidamente completados y protegidos.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 4,
        "Confirmación de Blueprint y creación coordinada de recursos en Render",
        EVID_DIR / "fig05-render-deploy-blueprint.png",
        "Captura de pantalla del resumen de recursos definidos en render.yaml (agrosmart-postgres y agrosmart-n8n) previo al inicio del aprovisionamiento.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 5,
        "Estado de sincronización Running durante el aprovisionamiento de recursos",
        EVID_DIR / "fig06-render-sync-running.png",
        "Captura de pantalla de la consola de Render reflejando la fase activa de sincronización y despliegue concurrente de la base de datos y el servicio web.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_heading_1(doc, "3. Diagnóstico de Incidencias Técnicas y Optimizaciones")

    add_heading_2(doc, "3.1. Corrección del registro de imagen Docker (PR #35)")
    add_paragraph_body(doc, "Durante el primer intento de despliegue, Render emitió el error 'the provided URL could not be fetched' al intentar resolver el dominio docker.n8n.io/n8nio/n8n:latest. Se solventó mediante el Pull Request #35, reorientando la ruta hacia el repositorio oficial en Docker Hub (docker.io/n8nio/n8n:1.83.2) e integrando la directiva N8N_ENCRYPTION_KEY con generación automática de clave.")

    add_heading_2(doc, "3.2. Optimización de memoria ante fallo OOM (Heap Crash)")
    add_paragraph_body(doc, "Debido a que el tier gratuito de Render cuenta con 512 MB de memoria RAM compartida, n8n sufrió un colapso de memoria durante el arranque (FATAL ERROR: JavaScript heap out of memory, exit code 134), provocando que la aplicación respondiera con Not Found y entrara en un ciclo de reinicios constantes.")

    add_figure_apa7(
        doc, 6,
        "Logs de arranque de n8n registrando inicialización en puerto 5678",
        EVID_DIR / "fig07-render-n8n-live-logs.png",
        "Captura de pantalla del registro de eventos de Render mostrando el mensaje n8n ready on port 5678 previo a la saturación de memoria en versiones no optimizadas.",
        "Elaboración propia a partir de los registros de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 7,
        "Respuesta Not Found en navegador por caída de servicio previa a estabilización",
        EVID_DIR / "fig08-n8n-not-found.png",
        "Captura de pantalla del navegador web reflejando la indisponibilidad temporal del servicio antes de aplicar las restricciones de consumo de memoria.",
        "Elaboración propia a partir de la navegación web."
    )

    add_paragraph_body(doc, "Para estabilizar el consumo de recursos se aplicaron las siguientes variables de entorno:")
    add_bullet_item(doc, "NODE_OPTIONS: --max-old-space-size=384 (fuerza la recolección periódica de basura en Node.js antes de sobrepasar el tope del contenedor de 512 MB).")
    add_bullet_item(doc, "N8N_RUNNERS_ENABLED: false (desactiva subprocesos concurrentes y procesadores auxiliares en segundo plano).")
    add_bullet_item(doc, "DB_POSTGRESDB_SSL_ENABLED: false (elimina la sobrecarga de negociación SSL en enlaces sobre la red privada interna).")
    add_bullet_item(doc, "PORT: 5678 y N8N_EDITOR_BASE_URL: https://agrosmart-n8n.onrender.com (alineación explícita de puertos y URLs públicas).")

    add_heading_2(doc, "3.3. Restablecimiento del esquema de base de datos (Error column User.role does not exist)")
    add_paragraph_body(doc, "Las interrupciones abruptas de n8n por falta de memoria provocaron que las migraciones internas de TypeORM quedaran en un estado corrupto e incompleto. Al acceder a la pantalla /setup para crear el usuario principal, el sistema arrojaba el error column User.role does not exist.")

    add_figure_apa7(
        doc, 8,
        "Panel de conexión interna y externa de agrosmart-postgres",
        EVID_DIR / "fig09-postgres-info-connect.png",
        "Captura de pantalla de la base de datos agrosmart-postgres en estado Available, indicando credenciales y parámetros de conexión interna y externa.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 9,
        "Configuración de reglas de firewall IP en PostgreSQL Inbound Rules",
        EVID_DIR / "fig10-postgres-networking-ip.png",
        "Captura de pantalla de la sección Networking de agrosmart-postgres permitiendo la conexión segura desde la dirección IP de la estación de trabajo local.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 10,
        "Error en asistente de creación de cuenta Owner: column User.role does not exist",
        EVID_DIR / "fig11-n8n-owner-error-bd.png",
        "Captura de pantalla del formulario de registro de n8n mostrando la notificación de error en base de datos debido al esquema corrupto.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_paragraph_body(doc, "El procedimiento de saneamiento consistió en: (a) habilitar temporalmente la IP local en las reglas de entrada de PostgreSQL, (b) conectarse mediante psql y ejecutar DROP SCHEMA public CASCADE; CREATE SCHEMA public; GRANT ALL ON SCHEMA public TO agrosmart;, y (c) ejecutar un Manual Deploy en n8n para que recreara limpiamente sus 25 migraciones estructurales sobre PostgreSQL 18.")

    add_heading_1(doc, "4. Validación Operativa, Migraciones y Carga Masiva de Datos")

    add_heading_2(doc, "4.1. Acceso a n8n y registro de cuenta Owner")
    add_paragraph_body(doc, "El servicio n8n arrancó de manera impecable en https://agrosmart-n8n.onrender.com. Se completó el asistente de incorporación creando la cuenta de propietario a nombre de Gabriel Emilio León Cangalaya.")

    add_figure_apa7(
        doc, 11,
        "Panel principal de n8n operativo con cuenta Owner autenticada",
        EVID_DIR / "fig12-n8n-dashboard-owner.png",
        "Captura de pantalla del panel principal de n8n confirmando la sesión activa del usuario Gabriel León y el espacio de trabajo preparado para flujos analíticos.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_heading_2(doc, "4.2. Ejecución de migraciones de negocio y carga masiva SQL")
    add_paragraph_body(doc, "Desde la consola PowerShell local y haciendo uso del cliente psql sobre la URL externa de agrosmart-postgres, se ejecutaron exitosamente los siguientes scripts:")
    add_bullet_item(doc, "001_schema.sql: Creación de la extensión vector, tablas de catálogo de productos, mercados, precios históricos y bitácora de auditoría.")
    add_bullet_item(doc, "002_indexes.sql: Generación de índices btree y compuestos para acelerar búsquedas analíticas por fechas y entidades.")
    add_bullet_item(doc, "002_seed_v2.sql: Inserción masiva de 9 445 registros oficiales de precios históricos del Gran Mercado Mayorista de Lima (GMML) mediante el comando COPY.")

    add_figure_apa7(
        doc, 12,
        "Ejecución exitosa de migraciones y seed de 9 445 filas en PostgreSQL Render",
        EVID_DIR / "fig14-psql-migraciones-ok.png",
        "Captura de pantalla de la terminal psql con las respuestas CREATE EXTENSION, CREATE TABLE, CREATE INDEX y COPY 9445 confirmando la persistencia de datos.",
        "Elaboración propia a partir de la consola de psql."
    )

    add_heading_2(doc, "4.3. Importación, compatibilidad y activación del flujo WF2")
    add_paragraph_body(doc, "Se importó en n8n el archivo oficial wf2-api-chat-analisis-predictivo.json. La primera importación reveló incompatibilidad de typeVersion: el JSON provenía de una versión más reciente que n8n 1.83.2 (Webhook 2.1, Gemini 1.1, AI Agent 3.1, Switch 3.4). n8n no reconocía esos nodos como ejecutables y al activar el flujo devolvía Cannot read properties of undefined (reading 'execute').")
    add_paragraph_body(doc, "Se normalizaron las versiones internas al conjunto soportado por 1.83.2, se completó el campo Data Name (informacion_agricola) en las herramientas de vector store y se sustituyó la clave de MapTiler embebida en el JSON por la expresión {{ $env.MAPTILER_API_KEY }}, en cumplimiento del principio de cero secretos en código.")

    add_figure_apa7(
        doc, 13,
        "Workflow WF2 importado en el editor visual de n8n",
        EVID_DIR / "fig13-wf2-workflow-importado.png",
        "Captura del lienzo de n8n con el grafo de nodos del flujo WF2 tras la importación inicial.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_figure_apa7(
        doc, 14,
        "Campo Data Name requerido en la herramienta Vector Store de WF2",
        EVID_DIR / "fig20-wf2-vectorstore-dataname.png",
        "Captura del nodo Answer questions with a vector store. El campo Data Name es obligatorio en typeVersion 1 y se completó con informacion_agricola.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_figure_apa7(
        doc, 15,
        "Flujo WF2 operativo en n8n cloud con Groq, Gemini, RAG y ramas SQL",
        EVID_DIR / "fig17-wf2-activo-canvas.png",
        "Captura del lienzo completo de WF2 tras la corrección de typeVersion y el enlace de credenciales Postgres, Groq, Gemini y Hugging Face.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_heading_2(doc, "4.4. Verificación pública del webhook de producción")
    add_paragraph_body(doc, "La URL de producción del webhook es https://agrosmart-n8n.onrender.com/webhook/v1/query. Una petición GET desde el navegador devolvió el JSON 404 «This webhook is not registered for GET requests. Did you mean to make a POST request?». Esa respuesta demuestra que el disparador está registrado y activo: n8n rechaza GET porque el nodo Webhook está configurado en POST.")

    add_figure_apa7(
        doc, 16,
        "Respuesta del webhook de producción ante una petición GET del navegador",
        EVID_DIR / "fig18-webhook-url-post.png",
        "Captura del navegador en https://agrosmart-n8n.onrender.com/webhook/v1/query. El código 404 aclara que el endpoint existe y espera POST, que es el método usado por n8nClient.js.",
        "Elaboración propia a partir de la navegación web."
    )

    add_heading_1(doc, "5. Conclusión")
    add_paragraph_body(doc, "La infraestructura de backend en Render quedó desplegada, optimizada y operativa. Se superaron los límites de memoria del plan gratuito, se saneó el esquema interno de n8n, se cargaron 9 445 registros del GMML, se compatibilizó WF2 con n8n 1.83.2 y se verificó el webhook público. El backend cloud del Sprint 1 queda listo para atender al frontend en Vercel.")

    add_references_section(doc, [
        "Equipo 2 AgroSmart Insights. (2026). Agrosmart-Insights [Repositorio de software]. GitHub. https://github.com/equipo2-AgroSmart-Insights/Agrosmart-Insights",
        "National Institute of Standards and Technology. (2022). Secure Software Development Framework (SSDF) Version 1.1 (NIST SP 800-218). https://csrc.nist.gov/Projects/ssdf",
        "n8n.io. (2026). Hosting n8n with PostgreSQL. Documentation. https://docs.n8n.io/hosting/databases/postgresql/",
        "Red Hat. (s. f.). ¿Qué es DevSecOps? https://www.redhat.com/es/topics/devops/what-is-devsecops",
        "Render. (2026). Blueprint specification. https://render.com/docs/blueprint-spec",
        "Render. (2026). Deploy n8n on Render. https://render.com/docs/deploy-n8n",
        "Render. (2026). PostgreSQL on Render. https://render.com/docs/databases",
    ])

    out_docx = OUT_DIR / "Informe-Configuracion-Render-Sprint1-V.1.docx"
    doc.save(out_docx)
    print(f"-> Guardado exitosamente: {out_docx}")
    return out_docx


# ==============================================================================
# 3. GENERACIÓN: INFORME 02 DE AVANCE SPRINT 1 (CONSOLIDADO DEL DÍA)
# ==============================================================================
def generate_avance_report():
    print("\n--- Generando Informe 02 de Avance Sprint 1 ---")
    doc = init_document()
    build_portada(doc, "INFORME DE AVANCE — SPRINT 1", subtitulo_informe="DESPLIEGUE CLOUD DEVSECOPS: VERCEL + RENDER")

    add_doc_title(doc, "INFORME DE AVANCE SPRINT 1 — DESPLIEGUE CLOUD DEVSECOPS")

    add_heading_1(doc, "1. Objetivo")
    add_paragraph_body(doc, "El presente informe tiene por objetivo documentar, sustentar y consolidar las actividades técnicas desarrolladas por el Líder DevSecOps durante las sesiones del 24 y 25 de agosto de 2026 del Sprint 1 en el proyecto AgroSmart Insights. El alcance comprende la integración controlada en main mediante Pull Requests, el despliegue del frontend en Vercel, el aprovisionamiento de n8n y PostgreSQL en Render, la resolución de incidencias de memoria y esquema, la carga de 9 445 registros del GMML, la activación de WF2 y la primera prueba de extremo a extremo en producción.")

    add_heading_2(doc, "1.1. Fundamento metodológico")
    add_paragraph_body(doc, "Se implementó un marco DevSecOps transversal articulado con el flujo de trabajo GitFlow institucional. Cada componente de software se somete a validaciones automáticas de seguridad, formateo estricto y pruebas unitarias previo a su fusión en la rama de producción, dando cumplimiento a las buenas prácticas de desarrollo seguro de software del estándar NIST SP 800-218 (NIST, 2022).")

    add_heading_2(doc, "1.2. Historias de usuario atendidas en la sesión")
    add_bullet_item(doc, "S1-01 — Desplegar servidor Cloud y unificar pipeline CI/CD en GitHub Actions (8 Story Points, Issue #28).")
    add_bullet_item(doc, "S1-06 — Configurar bloqueo estricto de PRs por fallos en pipeline de linters/tests (5 Story Points, Issue #29).")

    add_heading_1(doc, "2. Integración y Gobernanza del Repositorio (Pull Requests)")

    add_heading_2(doc, "2.1. Integración general del Sprint 1 (Pull Request #34)")
    add_paragraph_body(doc, "Se gestionó la integración del stack completo del Sprint 1 mediante el Pull Request #34 hacia main. Este PR incorporó el frontend React agrosmart-frontend, los flujos de n8n WF0/WF1/WF2, el esquema relacional con pgvector, el pipeline deploy.yml, la especificación render.yaml y los scripts operativos health-check.sh y rollback.sh. Los tres status checks obligatorios (validate-frontend, Auditar JSONs de n8n, Pruebas Unitarias de IA) finalizaron exitosamente y se obtuvo la aprobación del revisor Cristhian0520.")

    add_heading_2(doc, "2.2. Hotfix de infraestructura Blueprint (Pull Request #35)")
    add_paragraph_body(doc, "Durante el despliegue inicial en Render se identificó un fallo en la resolución de la imagen Docker de n8n. Se creó la rama fix/render-docker-image-url y se tramitó el PR #35, corrigiendo la URL hacia docker.io/n8nio/n8n:1.83.2 e incorporando N8N_ENCRYPTION_KEY con generación automática. El PR fue validado y fusionado, cerrando automáticamente los issues #28 y #29.")

    add_heading_1(doc, "3. Configuración y Despliegue en Vercel (Frontend)")
    add_paragraph_body(doc, "Se completó la puesta en producción del frontend agrosmart-frontend conectando el repositorio a Vercel bajo los siguientes parámetros técnicos:")
    add_bullet_item(doc, "Repositorio institucional: equipo2-AgroSmart-Insights/Agrosmart-Insights (rama main).")
    add_bullet_item(doc, "Root Directory: src/frontend/agrosmart-frontend (detección automática de Vite).")
    add_bullet_item(doc, "Variable de entorno: VITE_N8N_WEBHOOK_URL configurada con https://agrosmart-n8n.onrender.com/webhook/v1/query.")

    add_figure_apa7(
        doc, 1,
        "Importación y selección de Root Directory para el Frontend en Vercel",
        EVID_DIR / "fig01-vercel-importar-proyecto.png",
        "Captura de pantalla de la interfaz de Vercel vinculando el repositorio institucional y la rama de producción main.",
        "Elaboración propia a partir de la consola de Vercel (Vercel, 2026)."
    )

    add_figure_apa7(
        doc, 2,
        "Selector de Root Directory apuntando a src/frontend/agrosmart-frontend",
        EVID_DIR / "fig02-vercel-root-directory.png",
        "Captura de pantalla del selector de directorio en Vercel confirmando la ruta del microfrontend React y el framework Vite.",
        "Elaboración propia a partir de la consola de Vercel (Vercel, 2026)."
    )

    add_heading_1(doc, "4. Despliegue en Render (n8n + PostgreSQL)")

    add_heading_2(doc, "4.1. Creación de Blueprint Instance")
    add_paragraph_body(doc, "Se utilizó la funcionalidad Blueprint de Render basada en Infrastructure as Code (render.yaml). Al conectar el repositorio, Render analizó el archivo descriptivo e instanció de manera coordinada la base de datos gestionada agrosmart-postgres y el servicio web agrosmart-n8n.")

    add_figure_apa7(
        doc, 3,
        "Conexión del repositorio en Render para inicializar el Blueprint",
        EVID_DIR / "fig03-render-conectar-repo.png",
        "Captura de pantalla del catálogo de repositorios en Render mostrando la selección de Agrosmart-Insights.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 4,
        "Configuración de variables de entorno y API Keys en el Blueprint de Render",
        EVID_DIR / "fig04-render-blueprint-claves.png",
        "Captura de pantalla del formulario de Blueprint en Render con las variables del stack IA (Groq, Gemini, Hugging Face, MapTiler) ingresadas de forma segura.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 5,
        "Confirmación y creación de servicios agrosmart-postgres y agrosmart-n8n",
        EVID_DIR / "fig05-render-deploy-blueprint.png",
        "Captura de pantalla del resumen de recursos definidos en render.yaml antes de iniciar el aprovisionamiento en la nube.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 6,
        "Sincronización y despliegue simultáneo de recursos en Render",
        EVID_DIR / "fig06-render-sync-running.png",
        "Captura de pantalla mostrando el estado Running de la sincronización del Blueprint agrosmart-insights.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_heading_1(doc, "5. Diagnóstico y Resolución de Incidencias Técnicas")

    add_heading_2(doc, "5.1. Incidencia de memoria insuficiente (OOM) en instancia n8n")
    add_paragraph_body(doc, "Tras el despliegue inicial, el servicio agrosmart-n8n respondía con Not Found y se reiniciaba periódicamente. El análisis de logs reveló la excepción FATAL ERROR: JavaScript heap out of memory (código de salida 134), provocada por el consumo de memoria de Node.js al superar el límite de 512 MB del tier gratuito de Render.")

    add_figure_apa7(
        doc, 7,
        "Logs de Render mostrando arranque inicial y advertencias del servicio",
        EVID_DIR / "fig07-render-n8n-live-logs.png",
        "Captura de pantalla del visor de registros de Render con el mensaje n8n ready on port 5678 antes de registrar la saturación de heap.",
        "Elaboración propia a partir de los logs de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 8,
        "Fallo de disponibilidad (Not Found) en la URL pública de n8n",
        EVID_DIR / "fig08-n8n-not-found.png",
        "Captura de pantalla del navegador web reflejando la indisponibilidad temporal de la interfaz de n8n debido a los reinicios continuos por OOM.",
        "Elaboración propia a partir de la navegación web."
    )

    add_paragraph_body(doc, "Se resolvieron los fallos mediante las siguientes optimizaciones:")
    add_bullet_item(doc, "Fijación de versión: Se migró de :latest a la imagen estable docker.io/n8nio/n8n:1.83.2.")
    add_bullet_item(doc, "Límite estricto de Heap: Se inyectó NODE_OPTIONS=--max-old-space-size=384, restringiendo el recolector de basura de Node.js a 384 MB.")
    add_bullet_item(doc, "Desactivación de Runners: Se configuró N8N_RUNNERS_ENABLED=false para eliminar procesos auxiliares concurrentes.")
    add_bullet_item(doc, "Alineación de puerto y URL: Se definieron explícitamente PORT=5678 y N8N_EDITOR_BASE_URL.")

    add_heading_2(doc, "5.2. Incidencia de inconsistencia en esquema interno de n8n (User.role)")
    add_paragraph_body(doc, "Al acceder al asistente de bienvenida /setup, se produjo el error column User.role does not exist debido a que los reinicios por falta de memoria interrumpieron las migraciones internas automáticas de TypeORM en PostgreSQL.")

    add_figure_apa7(
        doc, 9,
        "Panel de conexión y parámetros de la base de datos agrosmart-postgres",
        EVID_DIR / "fig09-postgres-info-connect.png",
        "Captura de pantalla de la base de datos agrosmart-postgres en estado Available, detallando credenciales y puertos de acceso.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 10,
        "Habilitación de reglas de acceso IP en PostgreSQL Inbound Rules",
        EVID_DIR / "fig10-postgres-networking-ip.png",
        "Captura de pantalla de la sección Networking de la base de datos permitiendo el tráfico entrante desde la estación de trabajo del Líder DevSecOps.",
        "Elaboración propia a partir de la consola de Render (Render, 2026)."
    )

    add_figure_apa7(
        doc, 11,
        "Notificación de error en n8n: column User.role does not exist",
        EVID_DIR / "fig11-n8n-owner-error-bd.png",
        "Captura de pantalla del error al registrar la cuenta principal en /setup como consecuencia del esquema corrupto.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_paragraph_body(doc, "Para restablecer la integridad:")
    add_bullet_item(doc, "Se autorizó la IP pública de desarrollo en PostgreSQL Inbound IP Rules.")
    add_bullet_item(doc, "Se ejecutó mediante psql la sentencia DROP SCHEMA public CASCADE; CREATE SCHEMA public; GRANT ALL ON SCHEMA public TO agrosmart;.")
    add_bullet_item(doc, "Se reinició n8n con Manual Deploy, permitiendo que completara sus 25 migraciones limpias sobre PostgreSQL 18.")

    add_heading_1(doc, "6. Validación de Servicios Operativos y Carga de Datos")

    add_heading_2(doc, "6.1. Acceso a n8n y configuración de cuenta Owner")
    add_paragraph_body(doc, "El servicio n8n quedó 100% operativo en https://agrosmart-n8n.onrender.com. Se configuró exitosamente la cuenta administrativa a nombre de Gabriel Emilio León Cangalaya y se verificó el estado saludable del motor.")

    add_figure_apa7(
        doc, 12,
        "Dashboard principal de n8n con cuenta Owner configurada y activa",
        EVID_DIR / "fig12-n8n-dashboard-owner.png",
        "Captura de pantalla del panel operativo de n8n mostrando la sesión iniciada del usuario Gabriel León y el editor listo para ejecución de flujos.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_heading_2(doc, "6.2. Ejecución de migraciones de negocio y carga masiva SQL")
    add_paragraph_body(doc, "Desde la consola PowerShell local y mediante el cliente psql, se ejecutaron los scripts de base de datos sobre la instancia cloud de PostgreSQL:")
    add_bullet_item(doc, "001_schema.sql: Creación de extensiones vector, tablas productos, mercados, precios_historicos y bitácoras de auditoría.")
    add_bullet_item(doc, "002_indexes.sql: Creación de índices compuestos para optimización de consultas.")
    add_bullet_item(doc, "002_seed_v2.sql: Inserción masiva de 9 445 filas de precios históricos oficiales del Gran Mercado Mayorista de Lima (GMML).")

    add_figure_apa7(
        doc, 13,
        "Confirmación de ejecución de migraciones y carga de 9 445 registros en PostgreSQL",
        EVID_DIR / "fig14-psql-migraciones-ok.png",
        "Captura de pantalla de la terminal psql mostrando las respuestas CREATE EXTENSION, CREATE TABLE, CREATE INDEX y COPY 9445.",
        "Elaboración propia a partir de la consola de psql."
    )

    add_heading_2(doc, "6.3. Importación, compatibilidad y activación del flujo WF2")
    add_paragraph_body(doc, "Se importó wf2-api-chat-analisis-predictivo.json y se configuraron las credenciales Groq, Gemini, Hugging Face y PostgreSQL. El primer intento de activación falló por typeVersion posteriores a n8n 1.83.2 y por el campo Data Name vacío en las herramientas de vector store. Se corrigió el JSON en el repositorio, se reimportó el flujo y se verificó el webhook de producción.")

    add_figure_apa7(
        doc, 14,
        "Workflow WF2 importado en el editor visual de n8n",
        EVID_DIR / "fig13-wf2-workflow-importado.png",
        "Captura del lienzo de n8n con el grafo de nodos del flujo WF2 tras la importación.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_figure_apa7(
        doc, 15,
        "Flujo WF2 operativo con agentes Groq/Gemini, RAG y ramas SQL",
        EVID_DIR / "fig17-wf2-activo-canvas.png",
        "Captura del lienzo completo de WF2 ya enlazado a credenciales y listo para recibir POST en /webhook/v1/query.",
        "Elaboración propia a partir de la interfaz de n8n."
    )

    add_figure_apa7(
        doc, 16,
        "Confirmación pública de que el webhook WF2 está registrado y espera POST",
        EVID_DIR / "fig18-webhook-url-post.png",
        "Captura del navegador en la URL de producción. El 404 de método GET demuestra que el disparador está activo.",
        "Elaboración propia a partir de la navegación web."
    )

    add_heading_1(doc, "7. Despliegue del frontend y prueba de extremo a extremo")
    add_paragraph_body(doc, "El 25 de agosto de 2026 se publicó agrosmart-insights en Vercel (Hobby), con VITE_N8N_WEBHOOK_URL apuntando al webhook de Render. La primera consulta en el chat («¿Cuál es el precio actual de la papa?») evidenció que el cliente abortaba a los 6 segundos, mientras WF2 en Render continúa procesando. Se incrementó TIMEOUT_MS a 30 000 ms en n8nClient.js para el Pull Request de cierre.")

    add_figure_apa7(
        doc, 17,
        "Despliegue exitoso del frontend en Vercel",
        EVID_DIR / "fig16-vercel-deploy-ok.png",
        "Captura de la pantalla Congratulations de Vercel con la vista previa de AgroSmart Insights.",
        "Elaboración propia a partir de la consola de Vercel (Vercel, 2026)."
    )

    add_figure_apa7(
        doc, 18,
        "Timeout de 6 segundos en la primera prueba E2E del chat en Vercel",
        EVID_DIR / "fig19-frontend-timeout-6s.png",
        "Captura de la interfaz publicada. El mensaje en rojo corresponde al AbortController del cliente, no a una caída de n8n.",
        "Elaboración propia a partir de la aplicación desplegada en Vercel."
    )

    add_heading_1(doc, "8. Conclusión")
    add_paragraph_body(doc, "En las sesiones del 24 y 25 de agosto de 2026 se cerró el despliegue cloud del Sprint 1 bajo el rol de Líder DevSecOps: main protegida, frontend en Vercel, n8n y PostgreSQL en Render, 9 445 registros cargados, WF2 compatible con n8n 1.83.2 y webhook público verificado. El único cambio de código pendiente de fusión es el timeout del cliente NLQ (y la compatibilidad de typeVersion de WF2), que habilita la prueba E2E completa tras la aprobación del arquitecto.")

    add_references_section(doc, [
        "Equipo 2 AgroSmart Insights. (2026). Agrosmart-Insights [Repositorio de software]. GitHub. https://github.com/equipo2-AgroSmart-Insights/Agrosmart-Insights",
        "GitHub. (s. f.-a). Acerca de las ramas protegidas. Documentación de GitHub. https://docs.github.com/es/repositories/configuring-branches-and-merges-in-your-repository/defining-the-mergeability-of-pull-requests/about-protected-branches",
        "GitHub. (s. f.-b). Secretos en GitHub Actions. Documentación de GitHub. https://docs.github.com/es/actions/concepts/security-for-github-actions/security-guides/using-secrets-in-github-actions",
        "National Institute of Standards and Technology. (2022). Secure Software Development Framework (SSDF) Version 1.1 (NIST SP 800-218). https://csrc.nist.gov/Projects/ssdf",
        "Red Hat. (s. f.). ¿Qué es DevSecOps? https://www.redhat.com/es/topics/devops/what-is-devsecops",
        "Render. (2026). Blueprint specification. https://render.com/docs/blueprint-spec",
        "Render. (2026). Deploy n8n on Render. https://render.com/docs/deploy-n8n",
        "Vercel. (2026). Deploying a Git repository. https://vercel.com/docs/deployments/git",
        "Vercel. (2026). Monorepos. https://vercel.com/docs/monorepos",
    ])

    out_docx = OUT_DIR / "Informe-02-Avance-Sprint1-Cloud-DevSecOps-V.1.docx"
    doc.save(out_docx)
    print(f"-> Guardado exitosamente: {out_docx}")
    return out_docx


# ==============================================================================
# CONVERSIÓN COMPLETA A PDF
# ==============================================================================
def convert_all_to_pdf(docx_files):
    print("\n--- Convirtiendo todos los archivos a PDF con Word COM ---")
    wdFormatPDF = 17
    word = cc.CreateObject("Word.Application")
    word.Visible = False
    try:
        for docx_path in docx_files:
            src = pathlib.Path(docx_path)
            dst = src.with_suffix('.pdf')
            doc = word.Documents.Open(str(src))
            doc.SaveAs(str(dst), FileFormat=wdFormatPDF)
            doc.Close()
            print(f"PDF generado exitosamente: {dst.name} ({dst.stat().st_size:,} bytes)")
            time.sleep(1)
    finally:
        word.Quit()
    print("-> Conversion a PDF finalizada.")


def sync_to_repo(docx_files):
    print("\n--- Sincronizando documentos con el repositorio oficial ---")
    REPO_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    for docx_path in docx_files:
        src_docx = pathlib.Path(docx_path)
        src_pdf = src_docx.with_suffix('.pdf')
        
        dst_docx = REPO_DOCS_DIR / src_docx.name
        dst_pdf = REPO_DOCS_DIR / src_pdf.name
        
        shutil.copy2(src_docx, dst_docx)
        shutil.copy2(src_pdf, dst_pdf)
        print(f"Copiado a repo: {dst_docx.name} y {dst_pdf.name}")


if __name__ == '__main__':
    f1 = generate_vercel_report()
    f2 = generate_render_report()
    f3 = generate_avance_report()
    
    files = [f1, f2, f3]
    convert_all_to_pdf(files)
    sync_to_repo(files)
    print("\n=== PROCESO COMPLETADO AL 100% SIN ERRORES ===")
